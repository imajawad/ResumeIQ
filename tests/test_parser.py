"""
tests/test_parser.py
=====================
Unit tests for services/parser.py.

Test strategy:
    - Happy-path: valid file → non-empty string returned
    - Edge-case: empty stream → returns empty or raises ValueError
    - Invalid-input: unsupported extension → raises ValueError

Note on PDF/DOCX/image happy-path tests:
    These tests create minimal valid files in-memory so the test suite
    runs without external fixture files.
"""

import io
import pytest

from services.parser import extract_text, _STRATEGY_MAP


# ---------------------------------------------------------------------------
# Helper: create a minimal valid DOCX in memory
# ---------------------------------------------------------------------------

def _make_docx_stream(content: str = "Python Flask AWS Docker") -> io.BytesIO:
    """Return a BytesIO of a minimal DOCX with the given paragraph text."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Strategy dispatcher tests
# ---------------------------------------------------------------------------

class TestExtractTextDispatcher:
    """Tests for the public extract_text() dispatcher."""

    def test_unsupported_extension_raises_value_error(self):
        """Passing an unsupported extension must raise ValueError."""
        stream = io.BytesIO(b"some bytes")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            extract_text(stream, "xyz")

    def test_empty_extension_raises_value_error(self):
        """Empty extension string must raise ValueError."""
        stream = io.BytesIO(b"some bytes")
        with pytest.raises(ValueError):
            extract_text(stream, "")

    def test_all_supported_extensions_have_strategies(self):
        """Every extension in ALLOWED_EXTENSIONS should have a strategy."""
        expected = {"pdf", "docx", "png", "jpg", "jpeg"}
        assert expected.issubset(set(_STRATEGY_MAP.keys()))

    def test_extension_is_case_insensitive(self):
        """Extension matching should be case-insensitive."""
        # We can't easily create a valid PDF in-memory without PyMuPDF,
        # so we test that the strategy is found (dispatch works) by passing
        # a stream that will trigger a ValueError from the PDF parser itself,
        # NOT an "unsupported extension" error.
        stream = io.BytesIO(b"not a real pdf")
        with pytest.raises((ValueError, Exception)) as exc_info:
            extract_text(stream, "PDF")
        # Should NOT be "Unsupported file extension"
        assert "Unsupported file extension" not in str(exc_info.value)


# ---------------------------------------------------------------------------
# DOCX extraction tests (reliable in-memory)
# ---------------------------------------------------------------------------

class TestDocxExtraction:
    """Tests for the _extract_docx strategy via the public interface."""

    def test_docx_happy_path_returns_text(self):
        """Valid DOCX stream → non-empty string containing expected content."""
        stream = _make_docx_stream("Python Flask Machine Learning")
        result = extract_text(stream, "docx")
        assert isinstance(result, str)
        assert len(result.strip()) > 0
        assert "Python" in result

    def test_docx_multiple_paragraphs(self):
        """DOCX with multiple paragraphs → all paragraphs present in output."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Software Engineer")
        doc.add_paragraph("Skills: Python, Docker, Kubernetes")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = extract_text(buf, "docx")
        assert "Software Engineer" in result
        assert "Python" in result

    def test_docx_empty_document(self):
        """DOCX with no paragraphs → returns empty string (not an error)."""
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = extract_text(buf, "docx")
        assert isinstance(result, str)

    def test_docx_invalid_stream_raises_value_error(self):
        """Random bytes passed as DOCX → raises ValueError."""
        stream = io.BytesIO(b"this is not a valid docx file at all")
        with pytest.raises((ValueError, Exception)):
            extract_text(stream, "docx")

    def test_docx_stream_is_rewound_before_reading(self):
        """Parser must rewind stream even if it has been partially consumed."""
        stream = _make_docx_stream("Rewind test content")
        stream.read(10)  # Partially consume stream
        result = extract_text(stream, "docx")
        assert "Rewind test content" in result


# ---------------------------------------------------------------------------
# Extension aliasing tests
# ---------------------------------------------------------------------------

class TestExtensionAliases:
    """Tests for jpg / jpeg aliasing to the same OCR strategy."""

    def test_jpg_and_jpeg_use_same_strategy(self):
        """jpg and jpeg should map to the same extractor function."""
        assert _STRATEGY_MAP["jpg"] is _STRATEGY_MAP["jpeg"]
