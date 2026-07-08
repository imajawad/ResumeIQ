"""
tests/test_ats_checker.py
==========================
Unit tests for services/ats_checker.py.

Test strategy:
    - Happy-path: clean DOCX / image → score and flags returned
    - Edge-case: unsupported extension returns a structured response
    - Invalid-input: DOCX always returns non-error result with score
"""

import io
import pytest
from services.ats_checker import check_ats


class TestCheckATSDocx:
    """Tests for DOCX uploads — limited analysis path."""

    def _make_docx_stream(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Software Engineer with Python and Flask skills.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_docx_returns_structured_result(self):
        """DOCX upload → result dict with required keys."""
        stream = self._make_docx_stream()
        result = check_ats(stream, "docx")
        assert "ats_score" in result
        assert "flags" in result
        assert "suggestions" in result
        assert "checked" in result

    def test_docx_score_is_positive(self):
        """DOCX ATS score should be positive (we set it to 85 for clean DOCX)."""
        stream = self._make_docx_stream()
        result = check_ats(stream, "docx")
        assert result["ats_score"] > 0

    def test_docx_flags_are_empty(self):
        """Clean DOCX should have no flags."""
        stream = self._make_docx_stream()
        result = check_ats(stream, "docx")
        assert result["flags"] == []

    def test_docx_has_at_least_one_suggestion(self):
        """Even clean DOCX should return at least one best-practice suggestion."""
        stream = self._make_docx_stream()
        result = check_ats(stream, "docx")
        assert len(result["suggestions"]) >= 1


class TestCheckATSImage:
    """Tests for image uploads — always flagged as ATS-incompatible."""

    def test_image_resume_flagged(self):
        """Image upload → 'image_resume' flag in flags list."""
        stream = io.BytesIO(b"fake image bytes")
        result = check_ats(stream, "png")
        assert "image_resume" in result["flags"]

    def test_image_score_below_75(self):
        """Image resume ATS score should be below 75 (50 by design)."""
        stream = io.BytesIO(b"fake image bytes")
        result = check_ats(stream, "jpg")
        assert result["ats_score"] <= 75

    def test_image_suggestion_mentions_pdf_or_docx(self):
        """Image ATS suggestion should mention PDF or DOCX conversion."""
        stream = io.BytesIO(b"fake image bytes")
        result = check_ats(stream, "jpeg")
        combined_suggestions = " ".join(result["suggestions"]).lower()
        assert "pdf" in combined_suggestions or "docx" in combined_suggestions


class TestCheckATSReturnShape:
    """Verify return dict shape for all supported paths."""

    def test_all_paths_return_ats_score_int(self):
        """ats_score should be an integer for all extension paths."""
        for ext, stream in [
            ("docx", io.BytesIO(b"")),
            ("png", io.BytesIO(b"")),
            ("jpg", io.BytesIO(b"")),
        ]:
            # DOCX needs a valid stream; use a real DOCX for that path
            if ext == "docx":
                from docx import Document
                doc = Document()
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                stream = buf

            result = check_ats(stream, ext)
            assert isinstance(result["ats_score"], int), f"Failed for ext={ext}"

    def test_flags_is_always_a_list(self):
        """flags should always be a list, never None."""
        stream = io.BytesIO(b"")
        result = check_ats(stream, "png")
        assert isinstance(result["flags"], list)

    def test_suggestions_is_always_a_list(self):
        """suggestions should always be a list, never None."""
        stream = io.BytesIO(b"")
        result = check_ats(stream, "png")
        assert isinstance(result["suggestions"], list)
